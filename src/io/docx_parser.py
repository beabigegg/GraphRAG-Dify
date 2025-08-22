from __future__ import annotations
from typing import List
from docx import Document
from ..utils.chunking import pair_bilingual_lines, split_sections, split_with_overlap
from ..models.schema import TextChunk



def load_docx_text_chunks(path: str, rev: str | None = None) -> List[TextChunk]:
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    lines = pair_bilingual_lines(paras)

    # 改用滑動窗口切分段落（窗口大小 10 行，重疊 3 行）
    chunks = split_with_overlap(lines, rev=rev, window_size=10, overlap=3)
    # 如果仍希望使用原本的章節切分，可視情況二選一，不要先 return

    # 處理表格，每一行單獨成為一個 chunk
    for ti, table in enumerate(doc.tables):
        headers = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        for ri, row in enumerate(table.rows[1:], start=1):
            cells = [c.text.strip() for c in row.cells]
            chunks.append(
                TextChunk(
                    chunk_id=f"table-{ti}-row-{ri}",
                    section_id=f"table-{ti}",
                    rev=rev,
                    text=f"TABLE_ROW\nHEADERS: {headers}\nROW: {cells}",
                )
            )

    return chunks
